from docx import Document

# Create a new Document
doc = Document()

# Title
doc.add_heading('Mrwan Atef', 0)
doc.add_paragraph('Email: ma6655@fayoum.edu.eg')
doc.add_paragraph('LinkedIn: [Connect with me](your-linkedin-profile)')

# Objective
doc.add_heading('Objective', level=1)
doc.add_paragraph(
    'Passionate and skilled student of Mathematics, Computer Science, and Data Engineering, seeking to leverage my abilities in problem-solving, data analysis, and software development in a challenging Data Engineering role.'
)

# Education
doc.add_heading('Education', level=1)
doc.add_paragraph(
    'Bachelor of Science in Mathematics and Computer Science\n'
    'Fayoum University, Egypt\n'
    '(Expected Graduation: [Your Graduation Date])'
)

# Skills
doc.add_heading('Skills', level=1)

# Programming Languages
doc.add_heading('Programming Languages:', level=2)
doc.add_paragraph('- Python 🐍')
doc.add_paragraph('- SQL 🗄️')
doc.add_paragraph('- C++')

# Tools & Technologies
doc.add_heading('Tools & Technologies:', level=2)
doc.add_paragraph('- Git & GitHub 🗂️')

# Databases
doc.add_heading('Databases:', level=2)
doc.add_paragraph('- MySQL 🐬')
doc.add_paragraph('- PostgreSQL 🐘')
doc.add_paragraph('- MongoDB 🍃')

# Core Competencies
doc.add_heading('Core Competencies', level=1)
doc.add_paragraph('- Problem Solving: Ability to tackle complex and challenging problems with a logical and analytical approach.')
doc.add_paragraph('- Data Analysis: Skilled in extracting, cleaning, and analyzing data to derive meaningful insights.')
doc.add_paragraph('- Software Development: Experienced in writing clean, maintainable, and efficient code.')
doc.add_paragraph('- Collaboration: Effective communicator and team player who can coordinate well with cross-functional teams.')

# Experience
doc.add_heading('Experience', level=1)

# Open Source Contributor
doc.add_heading('Open Source Contributor', level=2)
doc.add_paragraph('[Project Name]')
doc.add_paragraph('[Date] - Present')
doc.add_paragraph('- Contributed to [Project Description].')
doc.add_paragraph('- Developed and optimized data pipelines using Python and SQL.')
doc.add_paragraph('- Implemented features and fixed bugs to improve project efficiency.')

# Data Engineering Intern
doc.add_heading('Data Engineering Intern', level=2)
doc.add_paragraph('[Company Name], [Location]')
doc.add_paragraph('[Date] - [Date]')
doc.add_paragraph('- Assisted in the design and development of scalable data infrastructure.')
doc.add_paragraph('- Supported the extraction, transformation, and loading (ETL) processes.')
doc.add_paragraph('- Collaborated with cross-functional teams to understand data requirements and deliver solutions.')

# Projects
doc.add_heading('Projects', level=1)

# Project 1
doc.add_heading('[Project Title]', level=2)
doc.add_paragraph('- Description: [Brief description of the project].')
doc.add_paragraph('- Technologies Used: Python, SQL, MySQL/PostgreSQL/MongoDB.')

# Project 2
doc.add_heading('[Project Title]', level=2)
doc.add_paragraph('- Description: [Brief description of the project].')
doc.add_paragraph('- Technologies Used: Python, Git & GitHub.')

# Currently Working On
doc.add_heading('Currently Working On', level=1)
doc.add_paragraph('- Enhancing my skills in machine learning and artificial intelligence.')
doc.add_paragraph('- Contributing to open-source projects in data engineering and analytics.')

# How to Reach Me
doc.add_heading('How to Reach Me', level=1)
doc.add_paragraph('Feel free to reach out if you’d like to collaborate on a project or just want to chat about tech and data!')

# References
doc.add_heading('References', level=1)
doc.add_paragraph('Available upon request.')

# Save the document
file_path = "/mnt/data/Mrwan_Atef_CV.docx"
doc.save(file_path)

file_path
